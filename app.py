import streamlit as st
import gspread
import pandas as pd
from google.oauth2.service_account import Credentials
import json
import io
from docx import Document

# --- 1. CONFIGURATION DE LA PAGE ---
st.set_page_config(page_title="GMAO Dashboard", layout="wide")

# --- FONCTION POUR GÉNÉRER LE RAPPORT WORD ---
def generate_word_report():
    doc = Document()
    doc.add_heading('Rapport de Projet : Système IoT et Tableau de Bord GMAO', 0)
    
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph("Ce projet vise à créer un système complet de Gestion de Maintenance Assistée par Ordinateur (GMAO) couplé à l'Internet des Objets (IoT). L'objectif est de collecter des données en temps réel depuis des machines industrielles à l'aide d'un microcontrôleur ESP32, de stocker ces données de manière sécurisée et accessible, et de les visualiser sur un tableau de bord interactif pour faciliter la prise de décision et la maintenance prédictive.")
    
    doc.add_heading('2. Architecture du Système', level=1)
    doc.add_paragraph("Le système repose sur une architecture \"Serverless\" (sans serveur backend dédié), privilégiant la simplicité, la légèreté et l'utilisation de services cloud gratuits et robustes.")
    doc.add_paragraph("L'architecture se décompose en quatre briques principales :")
    doc.add_paragraph("1. Acquisition des données (IoT) : Un microcontrôleur ESP32 simule la lecture de capteurs (température, statut de la machine).", style='List Number')
    doc.add_paragraph("2. Passerelle de communication (Webhook) : Un script Google Apps Script agit comme une API REST (Webhook) pour recevoir les données de l'ESP32 via des requêtes HTTP POST.", style='List Number')
    doc.add_paragraph("3. Stockage des données (Base de données) : Google Sheets est utilisé comme base de données légère pour stocker l'historique des relevés.", style='List Number')
    doc.add_paragraph("4. Visualisation et Analyse (Dashboard) : Une application web développée en Python avec le framework Streamlit lit les données depuis Google Sheets et affiche des indicateurs de performance (KPIs) et des graphiques en temps réel.", style='List Number')
    
    doc.add_heading('3. Composants Détaillés', level=1)
    doc.add_heading('3.1. Le Microcontrôleur ESP32 (IoT)', level=2)
    doc.add_paragraph("L'ESP32 est programmé en C++ (via l'IDE Arduino). Son rôle est de se connecter au réseau Wi-Fi local, acquérir les données des capteurs, formater ces données au format JSON, et envoyer ce payload vers l'URL du Webhook Google Apps Script.")
    
    doc.add_heading('3.2. Le Webhook (Google Apps Script)', level=2)
    doc.add_paragraph("Pour éviter à l'ESP32 de devoir gérer l'authentification complexe requise par l'API officielle de Google Sheets, nous utilisons Google Apps Script comme intermédiaire. Il intercepte les requêtes POST, parse le JSON, génère un horodatage et insère une nouvelle ligne dans le fichier Google Sheets.")
    
    doc.add_heading('3.3. La Base de Données (Google Sheets)', level=2)
    doc.add_paragraph("Google Sheets sert de base de données chronologique. Il est structuré avec les colonnes : Date_Heure, ID_machine, Température, Statut, Type_Panne, Durée_Intervention_min.")
    
    doc.add_heading('3.4. Le Tableau de Bord (Streamlit)', level=2)
    doc.add_paragraph("L'interface utilisateur est construite avec Streamlit. L'application utilise un Compte de Service Google Cloud pour l'authentification. Elle lit les données, calcule les KPIs (MTTR, total des pannes), affiche des graphiques et s'actualise toutes les 10 secondes.")
    
    doc.add_heading('4. Conclusion et Perspectives', level=1)
    doc.add_paragraph("Ce projet démontre la faisabilité de créer un système IoT/GMAO fonctionnel, peu coûteux et rapidement déployable. Améliorations futures possibles : utiliser de vrais capteurs industriels, ajouter un système d'alertes SMS/Email, ou migrer vers une base de données plus robuste comme Firebase.")
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- BARRE LATÉRALE (SIDEBAR) ---
st.sidebar.image("https://cdn-icons-png.flaticon.com/512/2043/2043074.png", width=100)
st.sidebar.header("📄 Rapport du Projet")
st.sidebar.write("Téléchargez le rapport complet du projet au format Word pour votre présentation.")

# Bouton de téléchargement Word
word_file = generate_word_report()
st.sidebar.download_button(
    label="📥 Télécharger le rapport (Word)",
    data=word_file,
    file_name="Rapport_Projet_GMAO.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# --- 2. CONNEXION SÉCURISÉE ---
scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
creds_dict = json.loads(st.secrets["google_credentials"])
creds = Credentials.from_service_account_info(creds_dict, scopes=scope)
client = gspread.authorize(creds)

# --- 3. LECTURE DES DONNÉES ---
@st.cache_data(ttl=10)
def load_data():
    sheet = client.open("GMAO_IoT_Data").sheet1
    data = sheet.get_all_records()
    df = pd.DataFrame(data)
    
    if not df.empty:
        df['Date_Heure'] = pd.to_datetime(df['Date_Heure'], format='mixed', dayfirst=True, errors='coerce')
        df['Durée_Intervention_min'] = pd.to_numeric(df['Durée_Intervention_min'], errors='coerce').fillna(0)
        df['Température'] = pd.to_numeric(df['Température'], errors='coerce').fillna(0)
        
    return df

# --- 4. AFFICHAGE DU DASHBOARD ---
st.title("⚙️ Dashboard GMAO & IoT")

try:
    df = load_data()
    
    if df.empty:
        st.warning("Le Google Sheet est vide pour le moment.")
    else:
        st.header("📊 Indicateurs de Performance (KPIs)")
        col1, col2, col3 = st.columns(3)
        
        last_temp = df.iloc[-1]['Température']
        last_status = df.iloc[-1]['Statut']
        temp_color = "inverse" if last_status == "Panne" else "normal"
        col1.metric("Température Actuelle", f"{last_temp} °C", last_status, delta_color=temp_color)
        
        interventions = df[df['Durée_Intervention_min'] > 0]
        mttr = interventions['Durée_Intervention_min'].mean() if not interventions.empty else 0
        col2.metric("MTTR (Temps moyen réparation)", f"{mttr:.1f} min")
        
        total_pannes = len(df[df['Statut'] == 'Panne'])
        col3.metric("Total Pannes Enregistrées", total_pannes)
        
        st.markdown("---")
        st.header("📈 Évolution de la Température")
        chart_data = df.dropna(subset=['Date_Heure']).set_index('Date_Heure')[['Température']]
        st.line_chart(chart_data)
        
        st.markdown("---")
        st.header("📋 Historique des Données Brutes")
        st.dataframe(df.iloc[::-1], use_container_width=True)
        
except Exception as e:
    st.error(f"Erreur lors de l'affichage : {e}")
